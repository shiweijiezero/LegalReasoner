import os
import json
from typing import Dict, List, Tuple
import re
from pathlib import Path
import win32com.client
from docx import Document as DocxDocument
import pythoncom
from tqdm import tqdm
from spire.doc import *
from spire.doc.common import *

class WordApplication:
    _instance = None

    @classmethod
    def get_instance(cls):
        """获取Word应用程序单例"""
        if cls._instance is None:
            pythoncom.CoInitialize()
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            cls._instance = word
        return cls._instance

    @classmethod
    def quit(cls):
        """关闭Word应用程序"""
        if cls._instance is not None:
            try:
                cls._instance.Quit()
            finally:
                cls._instance = None
                pythoncom.CoUninitialize()


class LegalDocumentAnalyzer:
    def __init__(self, content: str):
        self.content = content

    def extract_proceedings(self) -> List[Dict]:
        """Extract all proceedings from the document."""
        proceedings_pattern = r'([A-Z\s]+)NO\.?\s*(\d+)\s+OF\s+(\d+)'
        proceedings_matches = re.findall(proceedings_pattern, self.content)
        proceedings_matches = proceedings_matches[-1:]
        cases = []
        for case_type, case_no, year in proceedings_matches:
            case_pattern = fr'{case_type}NO\.?\s*{case_no}.*?(?=(?:[A-Z\s]+NO|$))'
            case_match = re.search(case_pattern, self.content, re.DOTALL)

            if case_match:
                case_section = case_match.group(0)
                case_type = case_type.strip()

                between_pattern = r'BETWEEN\s+(.*?)\s+(?:Plaintiff|Applicant|Petitioner)\s+and\s+(.*?)\s+(?:Defendant|Respondent)'
                parties_match = re.search(between_pattern, case_section, re.DOTALL)

                if parties_match:
                    plaintiff = parties_match.group(1).strip()
                    defendant = parties_match.group(2).strip()

                    cases.append({
                        "case_id": f"{case_type} {case_no}/{year}",
                        "type": case_type,
                        "plaintiff": plaintiff,
                        "defendant": defendant,
                        "raw_text": case_section[:200]
                    })

        return cases

    def extract_sections(self) -> Dict[str, str]:
        """
        Extract main sections from the document.
        Only returns sections if all required sections are found.
        """
        sections = {
            'claims': '',
            'facts': '',
            'issues': '',
            'disposal': ''
        }

        # Find section boundaries
        section_starts = {
            'claims': self.content.find('The claims'),
            'facts': self.content.find('The facts'),
            'issues': self.content.find('The issues'),
            'disposal': self.content.find('Disposal')
        }

        # Check if all required sections are present
        if any(pos == -1 for pos in section_starts.values()):
            # If any section is missing, return empty dict
            return {}

        # Create list of tuples (position, section_name) for sections
        valid_sections = [(pos, name) for name, pos in section_starts.items()]
        valid_sections.sort()  # Sort by position

        # Extract each section's content
        for i, (start_pos, section_name) in enumerate(valid_sections):
            end_pos = valid_sections[i + 1][0] if i < len(valid_sections) - 1 else len(self.content)
            sections[section_name] = self.content[start_pos:end_pos].strip()

        return sections

    def analyze(self) -> Dict:
        """
        Perform complete analysis of the document.
        Returns empty dict for sections if not all required sections are found.
        """
        return {
            "proceedings": self.extract_proceedings(),
            "sections": self.extract_sections()
        }


def convert_doc_to_docx(file_path: str) -> str:
    """
    Convert a .doc file to .docx format and return the path to the new file.
    """
    # Create absolute path
    print(f"Converting {file_path} to .docx format...")
    abs_path = os.path.abspath(file_path)
    # Create new file path
    new_file_path = f"{os.path.splitext(abs_path)[0]}.docx"

    # Check if converted file already exists
    if os.path.exists(new_file_path):
        return new_file_path

    # Get Word application instance
    word = WordApplication.get_instance()

    try:
        doc = word.Documents.Open(abs_path)
        doc.SaveAs(new_file_path, 12)  # 12 represents the Word format code for .docx
        doc.Close()
        return new_file_path
    except Exception as e:
        print(f"Error converting {file_path}: {str(e)}")
        raise

def read_doc(file_path: str) -> str:
    doc = Document()
    doc.LoadFromFile(file_path)
    content = doc.GetText()
    doc.Close()
    return content

def read_docx(file_path: str) -> str:
    """Read .docx file using python-docx."""
    doc = DocxDocument(file_path)
    return '\n'.join([paragraph.text for paragraph in doc.paragraphs])


def find_en_directories(base_path: str) -> List[Tuple[str, str]]:
    """
    Find all 'EN' directories and create corresponding output paths.
    Returns list of tuples (input_dir, output_dir).
    """
    input_output_pairs = []

    for root, dirs, _ in os.walk(base_path):
        if 'EN' in dirs:
            input_dir = os.path.join(root, 'EN')
            relative_path = os.path.relpath(input_dir, base_path)
            output_dir = os.path.join('processed_data', relative_path)
            input_output_pairs.append((input_dir, output_dir))

    return input_output_pairs


def process_directory(input_dir: str, output_dir: str):
    """
    Process all Word documents in the input directory and save results to output directory.
    Only save results if all required sections are present.
    """
    # Create output directory if it doesn't exist
    Path(output_dir).mkdir(parents=True, exist_ok=True)

    # Get absolute paths
    input_dir = os.path.abspath(input_dir)
    output_dir = os.path.abspath(output_dir)

    # Process each file
    for file_name in tqdm(os.listdir(input_dir), desc="Processing files"):
        if not (file_name.endswith('.doc') or file_name.endswith('.docx')):
            continue

        input_path = os.path.join(input_dir, file_name)
        output_path = os.path.join(
            output_dir,
            f"{os.path.splitext(file_name)[0]}.json"
        )

        try:
            # Convert .doc to .docx if necessary
            if file_name.endswith('.doc'):
                # 如果有同名.docx，那么就跳过
                if os.path.exists(f"{os.path.splitext(input_path)[0]}.docx"):
                    print(f"Skipped {input_path} - .docx file already exists")
                    continue
                # docx_path = convert_doc_to_docx(input_path)
                # content = read_docx(docx_path)
                content = read_doc(input_path)
            else:  # .docx
                content = read_docx(input_path)

            # Save content to .txt file
            with open(f"{os.path.splitext(input_path)[0]}.txt", 'w', encoding='utf-8') as f:
                f.write(content)

            # # Analyze document
            # analyzer = LegalDocumentAnalyzer(content)
            # analysis_results = analyzer.analyze()
            #
            # # Only save if all sections are present
            # if analysis_results["sections"]:
            #     with open(output_path, 'w', encoding='utf-8') as f:
            #         json.dump(analysis_results, f, ensure_ascii=False, indent=2)
            #     print(f"Processed and saved {input_path} -> {output_path}")
            # else:
            #     print(f"Skipped {input_path} - missing required sections")

        except Exception as e:
            print(f"Error processing {input_path}: {str(e)}")


def batch_convert_docs(directory: str):
    """
    Batch convert all .doc files in a directory to .docx format.
    """
    print(f"Converting .doc files in {directory} to .docx format...")
    for root, _, files in os.walk(directory):
        doc_files = [f for f in files if f.endswith('.doc')]
        for doc_file in tqdm(doc_files, desc="Converting files"):
            try:
                file_path = os.path.join(root, doc_file)
                convert_doc_to_docx(file_path)
            except Exception as e:
                print(f"Error converting {doc_file}: {str(e)}")


def main():
    try:
        base_path = "."

        # Find all EN directories and their corresponding output paths
        dir_pairs = find_en_directories(os.path.join(base_path, "data"))

        if not dir_pairs:
            print("No EN directories found!")
            return

        # First, convert all .doc files to .docx
        # for input_dir, _ in dir_pairs:
        #     batch_convert_docs(input_dir)

        # Then process each directory pair
        for input_dir, output_dir in dir_pairs:
            print(f"\nProcessing directory: {input_dir}")
            print(f"Output directory: {output_dir}")
            process_directory(input_dir, output_dir)

        print("\nAll processing complete!")

    finally:
        # 确保程序结束时关闭Word应用
        WordApplication.quit()


if __name__ == "__main__":
    main()